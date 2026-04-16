import re
import json
import asyncio
import uuid
from typing import List, Optional
from docx import Document
from pydantic import BaseModel
from openai import OpenAI


class ParsedRequirement(BaseModel):
    id: str
    title: str
    signal_interfaces: List[str] = []   # 信号接口名称列表
    scene_description: str = ''          # 场景描述
    function_description: str = ''       # 功能描述
    entry_condition: str = ''            # 功能触发条件
    execution_body: str = ''             # 功能进入后执行
    exit_condition: str = ''            # 功能退出条件
    post_exit_behavior: str = ''        # 功能退出后执行


# ---------- 文档文本提取 ----------

def extract_docx_text(file_path: str) -> str:
    """提取 Word 文档的全部文本内容（段落 + 表格），保留结构"""
    doc = Document(file_path)
    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = ' | '.join(cells)
            if line.strip('| '):
                parts.append(line)

    return '\n'.join(parts)


# ---------- LLM 需求解析 Prompt ----------

DOC_PARSE_SYSTEM_PROMPT = """你是一名汽车空调热管理系统的需求分析专家。你的任务是将需求文档原始文本解析为结构化的功能需求条目。

**一个完整需求的格式必须包含以下8个字段：**

1. **title** - 功能需求名称（如"蓝牙通话降风速"）
2. **signalInterfaces** - 信号接口列表，从文档中提取相关的输入/输出信号变量名（如 gCAN_BltCallSts、gCbnSys_xxx），支持后续 Excel 导入更新信息；若文档未提及则为空数组
3. **sceneDescription** - 场景描述，描述功能在什么工况下被触发
4. **functionDescription** - 功能描述，概括功能的整体行为
5. **entryCondition** - 功能触发条件，满足哪些条件时功能激活（如信号值、阈值、状态组合）
6. **executionBody** - 功能进入后执行，功能激活后系统执行的具体动作
7. **exitCondition** - 功能退出条件，满足哪些条件时功能退出
8. **postExitBehavior** - 功能退出后执行，功能退出后系统如何恢复或切换状态

**解析原则：**
- 章节编号（如 5.9.2）作为需求 ID，【重要】带小数点的子章节（如 5.9.2.1、5.9.2.2、5.9.2.3）是父章节的子内容，不是独立需求
- 示例结构：
  - 5.9.2 蓝牙通话降风速（父需求）
    - 5.9.2.1 信号接口
    - 5.9.2.2 场景描述
    - 5.9.2.3 功能描述（含 entry/exit/execution 等子项）
  - 【正确】5.9.2 解析为1个需求，包含所有8个字段
  - 【错误】把 5.9.2.1、5.9.2.2、5.9.2.3 各自解析为独立需求
- 章节标题（"5.9.2 蓝牙通话降风速"）提取为 title
- 功能描述/备注/注意/注等信息全部归入上述8个字段
- 【重要】"注:" 后面的补充说明内容必须归入对应的字段
- 保留文档中的所有信号名称、参数值（BlCallSts=0x1、FRZCU_PowerMode=0x1、温度阈值等）
- 只有带具体功能逻辑（entry/exit/execution等）的章节才算独立需求

**输出要求：**
- 每个顶级章节（如 5.9.2）输出为一个 JSON 对象，不是多个
- 不要把子章节（5.9.2.1、5.9.2.2、5.9.2.3）拆分为独立需求
- 输出格式示例：
```json
[
  {
    "id": "5.9.2",
    "title": "蓝牙通话降风速",
    "signalInterfaces": ["BltCallSts", "FRZCU_PowerMode"],
    "sceneDescription": "用户驾驶车辆进行蓝牙通话，空调自动降低风速。",
    "functionDescription": "TMS根据蓝牙通话状态控制降低鼓风机档位。",
    "entryCondition": "1、在空调开机状态下...",
    "executionBody": "当鼓风机风量≤6档时维持...",
    "exitCondition": "1、接收到蓝牙通话状态结束...",
    "postExitBehavior": "1、蓝牙通话结束后鼓风机风量恢复到原来状态..."
  }
]
```"""


DOC_PARSE_USER_TEMPLATE = """请解析以下需求文档，将其条目化为结构化需求列表：

---
{doc_text}
---

【重要】请务必：
1. title 字段必须是需求的功能名称，如"蓝牙通话降风速"，而不是"功能描述"这几个字
2. 每个顶级章节（如 5.9.2、5.10.1）解析为1个需求
3. "5.9.2" 是一级章节（总标题），其下的 "5.9.2.1"、"5.9.2.2" 是子内容，不是独立需求
4. title 应该从章节标题中提取（如"5.9.2 蓝牙通话降风速"中的"蓝牙通话降风速"）
5. 每个需求都要有完整的8个字段，特别是 title 不能为空或无意义文字

只输出 JSON 数组，不要输出其他内容。"""


# ---------- JSON 解析 ----------

def _parse_json_response(content: str) -> list[dict]:
    """从 LLM 响应中提取 JSON 数组"""
    # 直接解析
    try:
        result = json.loads(content)
        if isinstance(result, list):
            return result
    except json.JSONDecodeError:
        pass

    # markdown code block
    m = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', content, re.DOTALL)
    if m:
        try:
            result = json.loads(m.group(1))
            if isinstance(result, list):
                return result
        except json.JSONDecodeError:
            pass

    # 找 [ ] 包裹内容
    m = re.search(r'\[.*\]', content, re.DOTALL)
    if m:
        try:
            result = json.loads(m.group())
            if isinstance(result, list):
                return result
        except json.JSONDecodeError:
            pass

    raise ValueError(f'无法解析 LLM 响应为 JSON，原始内容: {content[:300]}')


# ---------- LLM 需求解析 ----------

async def parse_docx_with_llm(file_path: str, llm_config: dict) -> List[ParsedRequirement]:
    """使用 LLM 解析 Word 文档"""
    doc_text = extract_docx_text(file_path)

    if not doc_text.strip():
        raise ValueError('文档内容为空')

    user_prompt = DOC_PARSE_USER_TEMPLATE.format(doc_text=doc_text)

    def _call_llm():
        # MiniMax M2 系列需要特殊处理
        is_minimax = 'minimax' in llm_config['base_url'].lower() and 'MiniMax-M2' in llm_config['model']
        if is_minimax:
            import httpx
            url = f"{llm_config['base_url'].rstrip('/')}/v1/text/chatcompletion_v2"
            headers = {
                'Authorization': f'Bearer {llm_config["api_key"]}',
                'Content-Type': 'application/json',
            }
            if llm_config.get('group_id'):
                headers['GroupId'] = llm_config['group_id']
            with httpx.Client(timeout=120.0) as client:
                response = client.post(
                    url,
                    headers=headers,
                    json={
                        'model': llm_config['model'],
                        'messages': [
                            {'role': 'system', 'content': DOC_PARSE_SYSTEM_PROMPT},
                            {'role': 'user', 'content': user_prompt},
                        ],
                        'temperature': 0.3,
                        'max_tokens': 16384,
                    },
                )
            if response.status_code != 200:
                raise RuntimeError(f'MiniMax API error: {response.status_code} - {response.text}')
            resp_data = response.json()
            return resp_data.get('choices', [{}])[0].get('message', {}).get('content', '')
        else:
            client = OpenAI(
                api_key=llm_config['api_key'],
                base_url=llm_config['base_url'],
            )
            response = client.chat.completions.create(
                model=llm_config['model'],
                messages=[
                    {'role': 'system', 'content': DOC_PARSE_SYSTEM_PROMPT},
                    {'role': 'user', 'content': user_prompt},
                ],
                temperature=0.3,
                max_tokens=16384,
            )
            return response.choices[0].message.content or ''

    # 重试机制：失败后等待1秒，最多重试3次
    content = None
    last_error = None
    for attempt in range(3):
        try:
            content = await asyncio.to_thread(_call_llm)
            break
        except Exception as e:
            last_error = e
            if attempt < 2:
                import time
                time.sleep(1)
    if content is None:
        raise RuntimeError(f'LLM 调用失败（已重试3次）: {last_error}')
    parsed_raw = _parse_json_response(content)

    # 转换为 ParsedRequirement，确保必填字段
    requirements: List[ParsedRequirement] = []
    seen_ids: set = set()

    for item in parsed_raw:
        req_id = str(item.get('id', '')).strip()
        if not req_id:
            req_id = f"REQ-{len(requirements) + 1:03d}"

        # 去重
        original_id = req_id
        counter = 1
        while req_id in seen_ids:
            req_id = f"{original_id}-{counter}"
            counter += 1
        seen_ids.add(req_id)

        # 提取8个字段
        signal_interfaces = item.get('signalInterfaces', [])
        if isinstance(signal_interfaces, str):
            signal_interfaces = [s.strip() for s in signal_interfaces.split('\n') if s.strip()]

        req = ParsedRequirement(
            id=req_id,
            title=str(item.get('title', '')).strip() or req_id,
            signal_interfaces=signal_interfaces,
            scene_description=str(item.get('sceneDescription', '')).strip(),
            function_description=str(item.get('functionDescription', '')).strip(),
            entry_condition=str(item.get('entryCondition', '')).strip(),
            execution_body=str(item.get('executionBody', '')).strip(),
            exit_condition=str(item.get('exitCondition', '')).strip(),
            post_exit_behavior=str(item.get('postExitBehavior', '')).strip(),
        )
        requirements.append(req)

    return requirements
