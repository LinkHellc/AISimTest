from typing import List
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

HEADER_FONT = Font(name='微软雅黑', bold=True, size=10, color='FFFFFF')
HEADER_FILL = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
HEADER_ALIGNMENT = Alignment(horizontal='center', vertical='center', wrap_text=True)
CELL_ALIGNMENT = Alignment(vertical='top', wrap_text=True)
THIN_BORDER = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin'),
)

# 完整的 TestHarness 列（共15列）
EXCEL_HEADERS = [
    'TestModel',       # 0  测试模型名称（用户填写）
    'TestUnitModel',   # 1  需求/功能名称
    'TestHarnessName', # 2  用例名称 {req_id}_Tc_N
    'FreshFlag',       # 3  刷新标记，固定1
    'TestTime',        # 4  测试持续时间（秒）
    'TestStepName',    # 5  步骤名：Init/TSn/TVn
    'TestStepAction',  # 6  动作代码
    'TestTransition',  # 7  过渡条件 after(N,sec)
    'TestNextStepName',# 8  下一跳步骤名
    'TestVerifyName',  # 9  验证点名
    'WhenCondition',   # 10 触发条件时序
    'TestVerify',      # 11 验证Matlab脚本
    'TestDescription', # 12 测试步骤中文描述
    '优先级',          # 13（留空）
    '测试项目',        # 14（留空）
]


def export_to_excel(
    test_cases: List[dict],
    output_path: str,
    req_titles: dict[str, str] | None = None,
) -> str:
    """
    按 Matlab TestHarness/TestSequence 格式导出测试用例。
    每条用例结构：Init行 + TS行 + TV行 + TS行 + TV行 + ...（最后一个TS指向Init）

    列顺序（15列）：
    TestModel, TestUnitModel, TestHarnessName, FreshFlag, TestTime,
    TestStepName, TestStepAction, TestTransition, TestNextStepName,
    TestVerifyName, WhenCondition, TestVerify, TestDescription, 优先级, 测试项目

    Args:
        req_titles: 需求ID到标题的映射，用于 sheet 名称
    """
    req_titles = req_titles or {}
    wb = Workbook()
    # 移除默认 sheet
    wb.remove(wb.active)

    # 按 requirementId 分组
    cases_by_req: dict[str, list[dict]] = {}
    for tc in test_cases:
        rid = tc.get('requirementId', 'unknown')
        cases_by_req.setdefault(rid, []).append(tc)

    for req_id, cases in cases_by_req.items():
        # 生成 sheet 名称：取需求标题前31字符（Excel sheet 名限制）
        sheet_name = (req_titles.get(req_id) or req_id)[:31]
        # 避免 sheet 名称重复
        base_name = sheet_name
        idx = 1
        while sheet_name in wb.sheetnames:
            sheet_name = f"{base_name[:28]}_{idx}"
            idx += 1

        ws = wb.create_sheet(title=sheet_name)

        # 写入表头
        for col, header in enumerate(EXCEL_HEADERS, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = HEADER_ALIGNMENT
            cell.border = THIN_BORDER

        row_idx = 2
        for tc_idx, tc in enumerate(cases, 1):
            tc_name = tc.get('name', '')
            tc_id = tc.get('id', '')
            category = tc.get('category', 'positive')
            category_cn = '正例' if category == 'positive' else '反例'
            test_time = tc.get('testTime', 4)
            steps = tc.get('steps', [])

            harness_name = f"{req_id}_Tc_{tc_idx}"
            req_title_val = req_titles.get(req_id, tc.get('requirementId', ''))

            # 计算 TV 编号基数（用于 TV 行命名）
            # 每个 TS 配一个 TV，命名规则：TS1→TV1, TS2→TV2, ...
            # 结构：Init → TS1 → TV1 → TS2 → TV2 → ... → TSn → TVn

            if not steps:
                # 无步骤时只生成 Init 行
                ws.cell(row=row_idx, column=1, value='').border = THIN_BORDER          # TestModel
                ws.cell(row=row_idx, column=2, value=req_title_val).border = THIN_BORDER  # TestUnitModel
                ws.cell(row=row_idx, column=3, value=harness_name).border = THIN_BORDER  # TestHarnessName
                ws.cell(row=row_idx, column=4, value=1).border = THIN_BORDER            # FreshFlag
                ws.cell(row=row_idx, column=5, value=test_time).border = THIN_BORDER    # TestTime
                ws.cell(row=row_idx, column=6, value='Init').border = THIN_BORDER      # TestStepName
                ws.cell(row=row_idx, column=7, value='').border = THIN_BORDER          # TestStepAction
                ws.cell(row=row_idx, column=8, value='').border = THIN_BORDER          # TestTransition
                ws.cell(row=row_idx, column=9, value='TS1').border = THIN_BORDER       # TestNextStepName
                ws.cell(row=row_idx, column=10, value='').border = THIN_BORDER         # TestVerifyName
                ws.cell(row=row_idx, column=11, value='').border = THIN_BORDER        # WhenCondition
                ws.cell(row=row_idx, column=12, value='').border = THIN_BORDER        # TestVerify
                ws.cell(row=row_idx, column=13, value=f'{category_cn} - {tc_name}').border = THIN_BORDER  # TestDescription
                ws.cell(row=row_idx, column=14, value='').border = THIN_BORDER        # 优先级
                ws.cell(row=row_idx, column=15, value='').border = THIN_BORDER        # 测试项目
                for col in range(1, 16):
                    ws.cell(row=row_idx, column=col).alignment = CELL_ALIGNMENT
                row_idx += 1
                continue

            # Init 行（第1行）
            precondition = tc.get('precondition', '')
            ws.cell(row=row_idx, column=1, value='').border = THIN_BORDER               # TestModel
            ws.cell(row=row_idx, column=2, value=req_title_val).border = THIN_BORDER   # TestUnitModel
            ws.cell(row=row_idx, column=3, value=harness_name).border = THIN_BORDER    # TestHarnessName
            ws.cell(row=row_idx, column=4, value=1).border = THIN_BORDER              # FreshFlag
            ws.cell(row=row_idx, column=5, value=test_time).border = THIN_BORDER       # TestTime
            ws.cell(row=row_idx, column=6, value='Init').border = THIN_BORDER         # TestStepName
            ws.cell(row=row_idx, column=7, value=precondition).border = THIN_BORDER   # TestStepAction
            ws.cell(row=row_idx, column=8, value='').border = THIN_BORDER             # TestTransition
            ws.cell(row=row_idx, column=9, value='TS1').border = THIN_BORDER          # TestNextStepName
            ws.cell(row=row_idx, column=10, value='').border = THIN_BORDER           # TestVerifyName
            ws.cell(row=row_idx, column=11, value='').border = THIN_BORDER           # WhenCondition
            ws.cell(row=row_idx, column=12, value='').border = THIN_BORDER           # TestVerify
            ws.cell(row=row_idx, column=13, value=f'{category_cn} - {tc_name}').border = THIN_BORDER
            ws.cell(row=row_idx, column=14, value='').border = THIN_BORDER           # 优先级
            ws.cell(row=row_idx, column=15, value='').border = THIN_BORDER           # 测试项目
            for col in range(1, 16):
                ws.cell(row=row_idx, column=col).alignment = CELL_ALIGNMENT
            row_idx += 1

            # TS → TV 步骤对
            for i, step in enumerate(steps):
                ts_num = i + 1
                tv_num = i + 1

                # TS 行
                ts_action = step.get('TestStepAction', '')
                ts_transition = step.get('TestTransition', 'after(1,sec)')
                # 下一个TS或Init
                if i < len(steps) - 1:
                    ts_next = f'TS{ts_num + 1}'
                else:
                    ts_next = 'Init'
                tv_name = f'TV{tv_num}'
                when_cond = step.get('WhenCondition', 't>0.5 && t<4.5')
                ts_verify = step.get('TestVerify', '')
                ts_desc = step.get('TestDescription', '')

                ws.cell(row=row_idx, column=1, value='').border = THIN_BORDER        # TestModel
                ws.cell(row=row_idx, column=2, value='').border = THIN_BORDER        # TestUnitModel
                ws.cell(row=row_idx, column=3, value='').border = THIN_BORDER        # TestHarnessName
                ws.cell(row=row_idx, column=4, value=1).border = THIN_BORDER         # FreshFlag
                ws.cell(row=row_idx, column=5, value=test_time).border = THIN_BORDER  # TestTime
                ws.cell(row=row_idx, column=6, value=f'TS{ts_num}').border = THIN_BORDER   # TestStepName
                ws.cell(row=row_idx, column=7, value=ts_action).border = THIN_BORDER       # TestStepAction
                ws.cell(row=row_idx, column=8, value=ts_transition).border = THIN_BORDER   # TestTransition
                ws.cell(row=row_idx, column=9, value=ts_next).border = THIN_BORDER          # TestNextStepName
                ws.cell(row=row_idx, column=10, value=tv_name).border = THIN_BORDER        # TestVerifyName
                ws.cell(row=row_idx, column=11, value=when_cond).border = THIN_BORDER     # WhenCondition
                ws.cell(row=row_idx, column=12, value=ts_verify).border = THIN_BORDER     # TestVerify
                ws.cell(row=row_idx, column=13, value=ts_desc).border = THIN_BORDER       # TestDescription
                ws.cell(row=row_idx, column=14, value='').border = THIN_BORDER          # 优先级
                ws.cell(row=row_idx, column=15, value='').border = THIN_BORDER          # 测试项目
                for col in range(1, 16):
                    ws.cell(row=row_idx, column=col).alignment = CELL_ALIGNMENT
                row_idx += 1

                # TV 行（紧跟TS之后）
                ws.cell(row=row_idx, column=1, value='').border = THIN_BORDER   # TestModel
                ws.cell(row=row_idx, column=2, value='').border = THIN_BORDER   # TestUnitModel
                ws.cell(row=row_idx, column=3, value='').border = THIN_BORDER   # TestHarnessName
                ws.cell(row=row_idx, column=4, value=1).border = THIN_BORDER    # FreshFlag
                ws.cell(row=row_idx, column=5, value=test_time).border = THIN_BORDER  # TestTime
                ws.cell(row=row_idx, column=6, value=f'TV{tv_num}').border = THIN_BORDER   # TestStepName
                ws.cell(row=row_idx, column=7, value='').border = THIN_BORDER           # TestStepAction
                ws.cell(row=row_idx, column=8, value='').border = THIN_BORDER          # TestTransition
                ws.cell(row=row_idx, column=9, value='').border = THIN_BORDER          # TestNextStepName
                ws.cell(row=row_idx, column=10, value='').border = THIN_BORDER         # TestVerifyName
                ws.cell(row=row_idx, column=11, value='').border = THIN_BORDER         # WhenCondition
                ws.cell(row=row_idx, column=12, value='').border = THIN_BORDER        # TestVerify
                ws.cell(row=row_idx, column=13, value='').border = THIN_BORDER        # TestDescription
                ws.cell(row=row_idx, column=14, value='').border = THIN_BORDER        # 优先级
                ws.cell(row=row_idx, column=15, value='').border = THIN_BORDER        # 测试项目
                for col in range(1, 16):
                    ws.cell(row=row_idx, column=col).alignment = CELL_ALIGNMENT
                row_idx += 1

            # 空行分隔不同用例
            row_idx += 1

    # 列宽自适应
    col_widths = [30, 20, 25, 8, 8, 12, 55, 15, 12, 10, 18, 45, 30, 10, 10]
    for ws in wb.worksheets:
        for i, width in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width
        ws.freeze_panes = 'A2'

    wb.save(output_path)
    return output_path


def export_to_word(test_cases: List[dict], output_path: str) -> str:
    doc = Document()
    title = doc.add_heading('测试用例报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'共 {len(test_cases)} 条测试用例')
    positive_count = sum(1 for tc in test_cases if tc.get('category') == 'positive')
    doc.add_paragraph(f'正例: {positive_count} 条, 反例: {len(test_cases) - positive_count} 条')

    for tc in test_cases:
        doc.add_heading(f'{tc.get("id", "")} - {tc.get("name", "")}', level=2)
        table = doc.add_table(rows=5, cols=2, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cells_data = [
            ('类型', '正例' if tc.get('category') == 'positive' else '反例'),
            ('关联需求', tc.get('requirementId', '')),
            ('前提条件', tc.get('precondition', '')),
            ('测试步骤', '\n'.join(f'{i+1}. {s}' for i, s in enumerate(tc.get('steps', [])))),
            ('预期结果', tc.get('expectedResult', '')),
        ]
        for row_idx, (label, value) in enumerate(cells_data):
            table.cell(row_idx, 0).text = label
            table.cell(row_idx, 1).text = value
            for paragraph in table.cell(row_idx, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        doc.add_paragraph()

    doc.save(output_path)
    return output_path
